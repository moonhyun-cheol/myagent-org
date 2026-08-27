# -*- coding: utf-8 -*-
"""3-page CQR 컨셉 RA 발표용 Word — MY_prompt 위→아래 발표 순서."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

CONCEPT_ROOT = Path(__file__).resolve().parent.parent
OUT = CONCEPT_ROOT / "docs" / "CQR_컨셉RA_운영규칙_발표용.docx"

BODY = Pt(8)
HEAD1 = Pt(10)
HEAD2 = Pt(9)
TITLE = Pt(12)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = BODY
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEAD1
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEAD2
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_margins(doc)
    style_doc(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CQR 컨셉 RA — 운영 규칙 발표 (3장)")
    r.bold = True
    r.font.size = TITLE

    para(
        doc,
        "한 줄: CQR 촬영·상세·검수·개발 전담 AI. 매뉴얼(660줄)=일하는 법, "
        "회사 자료 16종=CQR·품목 아는 법. 평소 답=9칸 요약, 나머지는 요청 시만.",
    )

    # ── 1장: 시작 ──
    h1(doc, "【1장】 이 AI와 사용 방법")

    h2(doc, "CQR 컨셉 RA란")
    bullet(
        doc,
        "CQR 상품 촬영 컨셉·상세(A+) 이미지·업로드 검수·신제품 외관 스펙을 돕는 전용 AI",
    )
    bullet(
        doc, "CS·가격·재고 AI 아님. 없는 스펙(원단·재고 등)은 지어내지 않고 ‘확인 필요’"
    )

    h2(doc, "두 가지를 AI에 준다")
    bullet(
        doc,
        "① 업무 매뉴얼 MY_prompt.md (660줄) — 어떤 질문에 어떤 형식으로 답할지, 금지 사항",
    )
    bullet(
        doc,
        "② 회사 자료 16종 (전략서·품목표·촬영가이드 등) — CQR·모델코드·원단 내용",
    )
    bullet(
        doc,
        "통합본 MY_prompt_bundle.md (4,262줄) = ①+②를 한 파일로 합친 붙여넣기용",
    )

    h2(doc, "MY_prompt vs 통합본")
    bullet(doc, "MY_prompt 498~518줄: 16종 ‘이름·참고 순서’ 목록만 있음 → 본문 없음")
    bullet(doc, "통합본 796줄~: 전략서·품목표 등 본문이 # EMBEDDED 로 들어 있음")

    h2(doc, "용어 (아래 매뉴얼 읽을 때)")
    bullet(doc, "9칸 = 회의용 짧은 브리프(기본). 풀브리프(11칸) = 촬영팀용 전체 기획서")
    bullet(doc, "캐스팅 = 누가 모델로 나올지. 컷(컷시트) = 찍을 사진 한 장씩 목록")
    bullet(
        doc,
        ".art = AI 이미지용 영어 1장. QC = 업로드 검수표. .dev = 주머니·컬러 개발표",
    )
    bullet(doc, "운영모드(.ops) = 답 끝에 ‘다음에 뭘 할지’ 선택 메뉴 (평소엔 안 씀)")

    # ── 2장: 매뉴얼 상단~중반 (위→아래) ──
    h1(doc, "【2장】 매뉴얼 본문 — 위에서 아래로 (1~209줄)")

    h2(doc, "1~10줄 · 기본 약속")
    bullet(doc, "가격·재고·원단·사이즈 등 확인 안 된 정보 → 만들지 않음")
    bullet(doc, "내부 서버 경로·내부 코드명 사용자에게 노출 금지")

    h2(doc, "11~61줄 · 답 양식 (가장 중요)")
    bullet(
        doc,
        "평소 기본 → 9칸 (컨셉명·슬로건·무드·이미지 방향·라인·영화·CQR 연결 등 9항목)",
    )
    bullet(doc, "‘풀브리프/.ff/캐스팅/컷’ 요청 → 11칸+캐스팅+컷 목록 (촬영팀용)")
    bullet(
        doc, "‘.art/이미지 프롬프트’ 요청 → lifestyle 메인 1장 영어 (자동 제안 없음)"
    )
    bullet(
        doc,
        "‘.dev’ → 개발표 · ‘QC/검수’ → 검수표 · 요청한 것만, ‘이미지도 드릴까요?’ 금지",
    )

    h2(doc, "62~96줄 · 하는 일 / 안 하는 일")
    bullet(
        doc,
        "함: 촬영 브리프, 상세 방향, QC, 개발스pec / 안 함: CS, 가격, 관광·출근 모델",
    )
    bullet(doc, "순서: 질문 분류 → 품목표 조회 → 원단·장소 맞춤 → 캐스팅 → 양식 출력")

    h2(doc, "97~116줄 · 질문 자동 분류")
    bullet(doc, "촬영·컨셉→9칸 / 풀브리프·컷→11칸 / 이미지→영어 / QC→표 / 개발→표")
    bullet(doc, "모델코드·ASIN 있으면 품목표 조회 필수")

    h2(doc, "117~137줄 · CQR 브랜드 철학")
    bullet(
        doc, "목적이 스타일보다 앞선다 (PAA). 4라인: 리버레이터·커버트·새퍼·익스페디션"
    )
    bullet(doc, "장면: 누가·왜 / 무슨 일 / 어디서 — 전략서 내용을 답에 반영")

    h2(doc, "138~209줄 · 품질 기준")
    bullet(doc, "슬로건 매번 새 영문 12~18자 / 모델 11유형 로테이션")
    bullet(doc, "CQR 입은 사람 = 지금 하는 일 있어야 함 (관광·카페 노트북 금지)")
    bullet(doc, "AI 이미지 = ‘달라’고 할 때만 1장")

    # ── 3장: 매뉴얼 하단 (210~660줄) ──
    h1(doc, "【3장】 매뉴얼 본문 — 위에서 아래로 (210~660줄)")

    h2(doc, "210~391줄 · 말투·작업·운영모드")
    bullet(doc, "평소: 한국어 존댓말, 본문만 (메뉴 없음)")
    bullet(doc, "운영모드(.ops): 풀브리프/이미지/QC/개발 중 다음 작업 고르는 메뉴")
    bullet(doc, "촬영·이미지·QC·개발은 각각 다른 처리 흐름")

    h2(doc, "392~497줄 · 품목표·원단·장면")
    bullet(doc, "TLP125 등 모델코드 → 품목표에서 라인·원단 조회 (추측 금지)")
    bullet(
        doc, "경량 원단=여름·활동 장면 / 방한=추운 야외 (설원+여름용 같은 어색함 방지)"
    )

    h2(doc, "498~527줄 · 참조 문서 16종 (지식)")
    bullet(
        doc,
        "매뉴얼 여기엔 ‘이 16개를 참고하라’ 목록만 — 전략서·품목표 전문은 별도 파일",
    )
    bullet(
        doc,
        "전략서, 품목표, 개발방향, 컬러, 브랜드컨셉, 임무매뉴얼, 촬영가이드, "
        "비주얼DNA, QC가이드 등 → 통합본에 합치거나 knowledge로 붙임",
    )

    h2(doc, "528~618줄 · 양식 항목 상세")
    bullet(doc, "9칸: 9항목 채우고 CQR 연결에서 끝 (600~1200자)")
    bullet(
        doc,
        "11칸: 인물(캐스팅)·장소·임무·착장·촬영·영화·배우·컷시트까지 (촬영팀용)",
    )
    bullet(doc, "이미지: 짧은 브리프 + 메인 lifestyle 1장 영어")

    h2(doc, "619~636줄 · 임무 11종·라인별 금지")
    bullet(
        doc,
        "장면 분류: 이동·관측·통신·지휘·사냥견·현장검증·훈련·사냥정비·"
        "모토캠프·특수정찰·은밀정찰",
    )
    bullet(doc, "라인별 금지: 리버레이터 ‘교관’, 익스페디션 ‘등산’, 커버트 ‘요원’ 등")

    h2(doc, "637~660줄 · 우선순위·평소 기본값 (마무리)")
    bullet(doc, "충돌 시: 안전 → 허구금지 → 목적·임무 → 출력형식 → 말투")
    bullet(doc, "평소: 9칸 / 이미지·풀브리프·컷은 요청할 때만 / 자동 제안 없음")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 발표 끝 —")
    r.italic = True
    r.font.size = BODY

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
