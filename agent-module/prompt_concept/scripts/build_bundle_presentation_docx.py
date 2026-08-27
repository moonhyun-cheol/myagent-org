# -*- coding: utf-8 -*-
"""MY_prompt_bundle 발표용 Word 생성."""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "MY_prompt_bundle_발표용_뼈대.md"
OUT = ROOT / "docs" / "MY_prompt_bundle_발표용_뼈대.docx"


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.0)
        s.bottom_margin = Cm(1.0)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
    n = doc.styles["Normal"]
    n.font.name = "Malgun Gothic"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    n.font.size = Pt(9)

    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0 if "발표용" in line else 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif (
            line.startswith("|")
            or line.startswith("```")
            or line.startswith("┌")
            or line.startswith("│")
            or line.startswith("├")
            or line.startswith("└")
            or line.startswith("→")
        ):
            doc.add_paragraph(line)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("*"))
            r.bold = True
        elif line.strip() == "---":
            doc.add_paragraph("─" * 36)
        elif line.strip().startswith(">"):
            doc.add_paragraph(line.lstrip("> "))
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
