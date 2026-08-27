# -*- coding: utf-8 -*-
"""CQR 컨셉 RA 발표자료 — MY_prompt 핵심만 (슬라이드형 Word)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CONCEPT_ROOT = Path(__file__).resolve().parent.parent
OUT = CONCEPT_ROOT / "docs" / "CQR_컨셉RA_발표자료_핵심.docx"

TITLE = Pt(14)
SLIDE = Pt(11)
BODY = Pt(9)
SUB = Pt(8)


def setup(doc: Document) -> None:
    for s in doc.sections:
        s.top_margin = Cm(1.0)
        s.bottom_margin = Cm(1.0)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    n = doc.styles["Normal"]
    n.font.name = "Malgun Gothic"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    n.font.size = BODY
    pf = n.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(2)


def slide_title(doc: Document, n: int, title: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"슬라이드 {n}  |  {title}")
    r.bold = True
    r.font.size = SLIDE
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def line(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = BODY


def divider(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("─" * 40)
    r.font.size = SUB
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup(doc)

    # 표지
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CQR 컨셉 RA\n운영 규칙 — 발표 핵심")
    r.bold = True
    r.font.size = TITLE
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("MY_prompt 660줄 + 지식 16종 요약  |  대표님·비개발자용")
    r2.font.size = SUB

    # 1
    slide_title(doc, 1, "한 줄 요약")
    bullet(doc, "CQR 촬영·상세·검수·개발 전담 AI — CS·가격 AI 아님")
    bullet(doc, "매뉴얼(660줄)=일하는 법  +  회사자료 16종=CQR·품목 아는 법")
    bullet(doc, "평소 답=9칸 요약  |  풀브리프·이미지·QC·개발은 요청할 때만")
    bullet(doc, "없는 스펙 지어내지 않음  |  「이미지도 드릴까요?」 자동 제안 없음")

    # 2
    slide_title(doc, 2, "두 가지를 AI에 준다")
    line(doc, "① MY_prompt.md (660줄) — 어떤 질문에 어떤 형식으로 답할지, 금지 사항")
    line(doc, "② 회사 자료 16종 — 전략서·품목표·촬영가이드 등 (본문)")
    line(doc, "통합본 MY_prompt_bundle (4,262줄) = ①+② 한 파일 붙여넣기용")
    line(
        doc,
        "MY_prompt 498~518줄 = 16종 목록만 (본문 없음) → 통합본 796줄~에 본문 포함",
        bold=True,
    )

    # 3
    slide_title(doc, 3, "답 6가지 (기본은 9칸)")
    rows = [
        ("9칸 요약", "「컨셉 알려줘」", "평소 기본"),
        ("11칸 풀브리프", "「풀브리프 / .ff」", "촬영팀용"),
        ("AI 이미지", "「.art / 이미지 프롬프트」", "짧은 브리프+영어 1장"),
        ("QC", "「검수해줘」+ 업로드", "체크표"),
        ("개발표", "「.dev / 주머니 스펙」", "주머니·컬러 표"),
        ("라인 설명", "「CQR가 뭐야」", "가이드"),
    ]
    for a, b, c in rows:
        bullet(doc, f"{a}  ←  {b}  ({c})")

    # 4
    slide_title(doc, 4, "9칸 vs 11칸 (풀브리프)")
    line(doc, "9칸 = 기획 회의용 짧은 브리프 (약 600~1200자)")
    line(doc, "11칸 = 촬영장용 — 9칸 내용 + 아래가 추가됨:")
    bullet(doc, "캐스팅(누가) · 장소 · 시간·날씨 · 임무 · 착장 · 장비(로드아웃)")
    bullet(doc, "촬영 연출 · 매체 DNA · 배우 3티어 · 컷시트(찍을 사진 목록)")
    line(doc, "「풀브리프」라고 명시할 때만 11칸 — 평소 컨셉 질문은 9칸만")

    # 5
    slide_title(doc, 5, "용어 30초")
    terms = [
        ("브리프", "촬영·이미지 기획서"),
        ("캐스팅", "누가 모델로 나올지 (11유형 로테이션)"),
        ("컷/컷시트", "찍을 사진 한 장씩 목록"),
        (".art", "AI 이미지용 영어 1장 (lifestyle 메인)"),
        ("QC", "올린 이미지 브랜드 맞는지 검수표"),
        (".dev", "신제품 주머니·컬러 개발표"),
        ("운영모드 .ops", "다음 작업 고르는 메뉴 (평소엔 안 씀)"),
    ]
    for k, v in terms:
        bullet(doc, f"{k} = {v}")

    # 6
    slide_title(doc, 6, "매뉴얼 세 구간 — 역할 차이")
    bullet(doc, "11~61줄  답 양식 — 무슨 종류의 답을 줄지 + 기본 약속 (기본=9칸)")
    bullet(doc, "97~116줄  질문 자동 분류 — 사용자 말 → 9칸/11칸/.art/QC/.dev 중 선택")
    bullet(doc, "528~618줄  양식 항목 상세 — 고른 양식 안 항목을 어떻게 채울지")
    line(doc, "비유: 종류 고르기 → 접수 분류 → 양식지 작성", bold=True)

    # 7
    slide_title(doc, 7, "예시: 「TLP125 촬영 컨셉」이 거치는 줄")
    bullet(doc, "1~10  없는 정보 만들지 않음")
    bullet(doc, "11~61  9칸 확정, 이미지·풀브리프 자동 없음")
    bullet(doc, "97~116  9칸으로 분류 → TLP125 품목표 조회")
    bullet(doc, "117~209  CQR 철학·슬로건·임무 있는 인물")
    bullet(doc, "392~497  원단·장면 맞춤 (화면엔 TPO 헤더 안 보임)")
    bullet(doc, "528~531  9항목 순서대로 작성 → CQR 연결에서 끝")
    bullet(doc, "637~660  평소 기본값 재확인")
    line(doc, "안 거침: FULL/.art/.dev/QC 양식, 운영모드 메뉴, 컷시트")

    # 8
    slide_title(doc, 8, "지식 16종 — 한 줄 역할")
    klist = [
        "전략서 v3.1 — PAA·4라인·금지 (법)",
        "품목표 — TLP125→라인·원단",
        "개발방향 — 모델별 배경·컨셉 메모",
        "브랜드 컨셉 — 라인·캐릭터 설명 (전략서 요약판)",
        "로드아웃 — 임무 11종·장비 규칙",
        "촬영 브리프 엔진 — 풀브리프 작성법",
        "비주얼 DNA — 상세 이미지·TPO·QC",
        "AI 이미지 조립법 — .art 영어 만드는 법",
        "플레이북 — 업무 SOP (분류·QC·.ops)",
        "영화 레퍼런스 — 풀브리프 📽️용",
        "카탈로그 — ASIN·listing 연결 (창작용 아님)",
    ]
    for k in klist:
        bullet(doc, k)

    # 9
    slide_title(doc, 9, "전략서 vs 브랜드 컨셉 vs 카탈로그")
    bullet(
        doc,
        "내부 전략서 v3.1 — 규칙·금지·임무의 원본 (data/CQR_INTERNAL_STRATEGY_v3.1.md)",
    )
    bullet(doc, "브랜드 컨셉 — 기획용 설명서, 충돌 시 전략서 우선")
    bullet(doc, "카탈로그 — ASIN·상품명 연결, listing 스냅샷 (컨셉 창작용 아님)")
    line(doc, "라인별 금지(교관·등산·요원·노동자) 출처 = 전략서 v3.1 → 매뉴얼 C-5b")

    # 10
    slide_title(doc, 10, "규칙이 겹치는 이유")
    bullet(doc, "중요 규칙은 AI가 놓치지 않게 앞·중간·뒤에 반복")
    bullet(doc, "충돌 시: 안전·법 → 허구금지 → 목적·임무 → 출력형식 → 말투")
    bullet(doc, "「안전」= 아마존·법·폭력·유명인 사칭 등 최상위 선")

    # 11
    slide_title(doc, 11, "실무 사용")
    bullet(doc, "Gemini 등: 통합본(MY_prompt_bundle) 붙여넣기")
    bullet(doc, "또는 MY_prompt + 16종 knowledge 각각 첨부")
    bullet(doc, "기획서 final_product_plan.md 는 16종에 없음 → 채팅에 따로 첨부")
    bullet(
        doc,
        "품목표·전략서 갱신 후: python prompt_concept/scripts/build_local_bundle.py",
    )

    divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—  발표 끝  —")
    r.italic = True
    r.font.size = SUB

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
