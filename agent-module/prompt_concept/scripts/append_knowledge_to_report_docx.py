# -*- coding: utf-8 -*-
"""Append knowledge-section pages to existing CQR 발표용 docx (does not overwrite user edits)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

CONCEPT_ROOT = Path(__file__).resolve().parent.parent
DOCX = CONCEPT_ROOT / "docs" / "CQR_컨셉RA_운영규칙_발표용.docx"

BODY = Pt(8)
HEAD1 = Pt(10)
HEAD2 = Pt(9)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEAD1
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEAD2
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


EMBEDDED = [
    (
        "796줄~",
        "품목표 (MODEL ROW INDEX)",
        "TLP125 등 모델코드→라인·원단·계절. AI가 추측 대신 여기서 조회",
    ),
    ("915줄~", "개발스펙 엔진", "신제품 .dev 요청 시 주머니·허리·컬러 표 작성 규칙"),
    ("1076줄~", "컬러코드", "색상 코드→색상명"),
    ("1103줄~", "개발방향", "신제품·라인 개발 방향 메모"),
    ("1716줄~", "카탈로그 요약", "ASIN·상품명 연결"),
    ("1782줄~", "브랜드 컨셉", "CQR 브랜드 정의·톤"),
    ("2031줄~", "내부 전략서 v3.1", "PAA·4라인·임무 11종·카피·비주얼 금지 — 답의 뼈대"),
    ("2275줄~", "로드아웃(임무) 매뉴얼", "이동·관측·통신… 11종 임무별 장비·장면 규칙"),
    ("2676줄~", "촬영 브리프 엔진", "풀브리프·9칸 작성 상세 가이드"),
    ("2878줄~", "슬로건 가이드", "영문 슬로건 12~18자 작성법"),
    ("2975줄~", "모델 캐스팅", "11가지 남성 유형·로테이션"),
    ("3231줄~", "비주얼 DNA", "상세페이지 이미지 슬롯·구도 분석"),
    ("3753줄~", "AI 이미지 조립법", ".art 요청 시 영어 프롬프트 만드는 방법"),
    ("3986줄~", "브랜드 이미지 플레이북", "컨셉 RA 역할·QC·운영모드 요약"),
    ("4129줄~", "영화 무드 참고", "풀브리프 📽️ 항목용 영화 레퍼런스"),
    ("4221줄~", "양식 예시 (GOLDEN)", "9칸·11칸 출력 형식 참고용"),
]


def append_knowledge_section(doc: Document) -> None:
    doc.add_page_break()

    h1(doc, "【부록】통합본(MY_prompt_bundle) 뒤쪽 — 지식(임베드) 부분")

    para(
        doc,
        "MY_prompt.md(660줄)는 ‘일하는 규칙’만 있습니다. "
        "통합본(약 4,262줄)은 792줄까지 규칙 + 그 아래(796줄~)에 "
        "회사 자료 16종 본문이 # EMBEDDED: … 제목으로 붙어 있습니다.",
    )
    para(
        doc,
        "Gemini 등에 통합본을 통째로 붙이면, AI가 규칙과 지식을 "
        "한 번에 읽습니다. MY_prompt만 쓸 때는 16종을 knowledge로 따로 붙여야 합니다.",
    )

    h2(doc, "구조 (위→아래)")
    bullet(doc, "1~792줄: 매뉴얼(규칙) — 앞 3장에서 설명한 내용")
    bullet(doc, "794줄: 구분선 (---)")
    bullet(doc, "796~4,262줄: 지식 본문 16블록 (아래 표)")

    h2(doc, "AI가 참고하는 순서 (630줄 목록 — 검색 우선순위)")
    bullet(
        doc, "1 전략서 → 2 품목표 → 3 개발스펙엔진 → 4 컬러 → 5 개발방향 → 6 브랜드컨셉"
    )
    bullet(
        doc, "7 로드아웃 → 8 촬영브리프엔진 → 9 슬로건 → 10 모델캐스팅 → 11 비주얼DNA"
    )
    bullet(
        doc, "12 AI이미지조립법 → 13 플레이북 → 14 영화무드 → 15 양식예시 → 16 카탈로그"
    )
    bullet(doc, "충돌 시: 사용자가 채팅에 붙인 최신 자료 → 품목표 → 개발방향")

    h2(doc, "16종 지식 — 파일에 붙은 순서·역할 (796줄~)")

    for loc, name, desc in EMBEDDED:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{loc} {name}: ")
        r1.bold = True
        p.add_run(desc)

    h2(doc, "실무에서 기억할 것")
    bullet(doc, "9칸·풀브리프 ‘형식’은 매뉴얼(앞부분)이 정함")
    bullet(doc, "‘TLP125가 어느 라인·원단인지’ 같은 내용은 품목표(796줄~)에서 가져옴")
    bullet(
        doc, "‘리버레이터 톤·임무 11종’은 전략서(2031줄~)·로드아웃(2275줄~)에서 가져옴"
    )
    bullet(
        doc,
        "품목표·전략서 갱신 후 → python scripts/build_local_bundle.py 로 통합본 재생성",
    )

    h2(doc, "MY_prompt vs 통합본 — 지식만 비교")
    bullet(doc, "MY_prompt 498~518줄: 16종 이름·순서만 (본문 없음)")
    bullet(doc, "통합본 796줄~: 16종 본문 전부 포함 (약 3,500줄)")
    bullet(doc, "기획서 final_product_plan.md 는 16종에 없음 → 채팅에 따로 첨부")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 지식 부분 설명 끝 —")
    r.italic = True
    r.font.size = BODY


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(
            f"파일 없음: {DOCX}\n먼저 build_concept_ra_report_docx.py 를 실행하세요."
        )

    doc = Document(str(DOCX))
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = BODY
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(2)

    append_knowledge_section(doc)
    doc.save(str(DOCX))
    print(f"Appended knowledge section to: {DOCX}")


if __name__ == "__main__":
    main()
